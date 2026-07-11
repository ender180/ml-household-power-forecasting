from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


RESULTS_DIR = Path("results_weather")
ASSET_DIR = Path("reports/academic_assets")
REPORT_DIR = Path("reports")
DOCX_PATH = REPORT_DIR / "机器学习课程考核报告_学术版_家庭电力消耗预测.docx"
MD_PATH = REPORT_DIR / "机器学习课程考核报告_学术版_家庭电力消耗预测.md"

MODEL_ORDER = ["lstm", "transformer", "conv_transformer"]
MODEL_LABELS = {
    "lstm": "LSTM",
    "transformer": "Transformer",
    "conv_transformer": "多尺度卷积 Transformer",
}


def set_run_font(run, size: float | None = None, bold: bool | None = None, east_asia: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), "Times New Roman")
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_heading_run_font(run, size: float, bold: bool = True) -> None:
    """Use one explicit font for heading numbers and Chinese text to avoid PDF font fallback."""
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key in ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]:
        r_fonts.set(qn(key), "Microsoft YaHei")
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def configure_styles(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.95)
    sec.right_margin = Inches(0.95)

    for style_name in ["Normal", "Body Text"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.22
        style.paragraph_format.space_after = Pt(3)

    for style_name, size, before, after in [
        ("Heading 1", 15, 10, 5),
        ("Heading 2", 12.2, 6, 3),
        ("Heading 3", 11, 4, 2),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        for key in ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]:
            style._element.rPr.rFonts.set(qn(key), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(2)


def heading(doc: Document, text: str, level: int):
    para = doc.add_paragraph(style=f"Heading {level}")
    size = {1: 15, 2: 12.2, 3: 11}.get(level, 11)
    run = para.add_run(text)
    set_heading_run_font(run, size)
    return para


def p(doc: Document, text: str, first_line: bool = True, after: float = 3.0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para = doc.add_paragraph()
    para.alignment = align
    if first_line:
        para.paragraph_format.first_line_indent = Inches(0.22)
    para.paragraph_format.line_spacing = 1.22
    para.paragraph_format.space_after = Pt(after)
    run = para.add_run(text)
    set_run_font(run, 10.5)
    return para


def centered(doc: Document, text: str, size: float = 10.5, bold: bool = False):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    set_run_font(run, size, bold)
    return para


def caption(doc: Document, text: str):
    para = centered(doc, text, 9, False)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(5)


def add_picture(doc: Document, path: Path, cap: str, width: float = 6.05) -> None:
    if not path.exists():
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    para.add_run().add_picture(str(path), width=Inches(width))
    caption(doc, cap)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, spec in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        if spec is None:
            element.set(qn("w:val"), "nil")
        else:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), str(spec.get("sz", 8)))
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")


def set_cell_margins(cell, top=70, start=95, bottom=70, end=95) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    set_run_font(run, 9.2, bold)


def academic_table(doc: Document, title: str, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    centered(doc, title, 9.4, True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    if widths:
        for i, width in enumerate(widths):
            table.columns[i].width = Inches(width)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_border(cell, top={"sz": 12}, bottom={"sz": 8})
            elif r_idx == len(table.rows) - 1:
                set_cell_border(cell, bottom={"sz": 12})
            else:
                set_cell_border(cell)
    doc.add_paragraph()


def add_pseudocode(doc: Document, title: str, lines: list[str]) -> None:
    centered(doc, title, 9.4, True)
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_border(cell, top={"sz": 8}, bottom={"sz": 8}, left={"sz": 8}, right={"sz": 8})
    set_cell_margins(cell, top=95, bottom=95, start=145, end=145)
    cell.text = ""
    for idx, line in enumerate(lines):
        para = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(8.8)
        run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()


def fmt(mean: float, std: float) -> str:
    return f"{mean:.2f} ± {std:.2f}"


def markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def build_rows(summary: pd.DataFrame, best: pd.DataFrame) -> tuple[list[list[str]], list[list[str]]]:
    ordered = summary.set_index(["model", "horizon"]).loc[
        [(m, h) for h in [90, 365] for m in MODEL_ORDER]
    ].reset_index()
    summary_rows = [
        [
            MODEL_LABELS[r["model"]],
            str(int(r["horizon"])),
            fmt(r["mse_mean"], r["mse_std"]),
            fmt(r["mae_mean"], r["mae_std"]),
            f"{r['val_loss_mean']:.4f}",
            str(int(r["runs"])),
        ]
        for _, r in ordered.iterrows()
    ]

    ordered_best = best.set_index(["model", "horizon"]).loc[
        [(m, h) for h in [90, 365] for m in MODEL_ORDER]
    ].reset_index()
    best_rows = [
        [
            MODEL_LABELS[r["model"]],
            str(int(r["horizon"])),
            str(int(r["seed"])),
            f"{r['test_mse']:.2f}",
            f"{r['test_mae']:.2f}",
        ]
        for _, r in ordered_best.iterrows()
    ]
    return summary_rows, best_rows


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    summary = pd.read_csv(RESULTS_DIR / "summary.csv")
    best = pd.read_csv(ASSET_DIR / "best_plot_seeds.csv")
    summary_rows, best_rows = build_rows(summary, best)

    doc = Document()
    configure_styles(doc)

    centered(doc, "家庭电力消耗多变量时间序列预测研究", 18, True)
    centered(doc, "作者：请填写姓名、学号、研究方向", 10, False)
    centered(doc, "代码仓库：请填写 GitHub 链接", 10, False)

    heading(doc, "1. 问题介绍", level=1)
    p(doc, "家庭电力消耗预测是智能家居能源管理与电网负荷调度中的重要问题。准确预测家庭未来用电曲线，不仅有助于居民理解自身用电行为、优化用电计划和降低电费，也能够为电力公司进行需求响应、峰谷调节与分布式能源管理提供参考。由于家庭用电同时受到季节变化、天气条件、设备使用习惯和家庭成员行为模式等因素影响，该任务具有明显的非线性、非平稳性与多变量耦合特征。")
    p(doc, "本文以 UCI Individual household electric power consumption 数据集为基础，构建基于过去 90 天多变量序列预测未来每日总有功功率的实验流程。原始用电数据为分钟级记录，直接建模会带来较高计算成本和较强噪声，因此本文按照日粒度进行汇总，并将 Météo-France 月度基础气候数据作为外部天气特征融合到样本中。最终任务包括 90 天短期预测和 365 天长期预测，两种预测长度分别训练和评估。")
    p(doc, "用电数据包含 global_active_power、global_reactive_power、voltage、global_intensity 和三个分表能耗变量。预处理时，global_active_power、global_reactive_power 与 sub_metering_1/2/3 按天求和，voltage 与 global_intensity 按天求平均，并计算 sub_metering_remainder 作为分表未覆盖的剩余能耗。天气特征选取 RR、NBJRR1、NBJRR5、NBJRR10 和 NBJBROU，预处理脚本在巴黎及周边省份气象站中选择字段完整且覆盖 2006-2010 年的 PARIS-MONTSOURIS 站点，并按月份合并到每日样本。")
    p(doc, "样本构造采用时间顺序滑动窗口方式：对每一个可用日期 t，模型读取其之前连续 90 天的多变量特征，并预测 t 之后 H 天的每日有功功率序列。数据划分严格按时间先后进行，测试集保留最后 455 天，以保证 365 天预测任务仍具有完整的 90 天历史上下文。该设置避免了随机划分在时间序列任务中可能造成的未来信息泄漏。")
    p(doc, "从任务性质看，短期预测主要考察模型对近期负荷水平、局部波动和短周期行为的刻画能力；长期预测则更依赖对趋势、季节变化和周期模式的综合建模。因而，本文不仅比较传统循环结构和自注意力结构，也进一步引入局部多尺度卷积特征，以观察显式局部模式提取对长期预测的作用。")
    academic_table(
        doc,
        "表 1  主要输入变量及处理方式",
        ["变量类别", "变量", "日级处理方式"],
        [
            ["目标变量", "global_active_power", "按天求和，作为预测目标"],
            ["用电变量", "global_reactive_power, sub_metering_1/2/3", "按天求和"],
            ["电气状态", "voltage, global_intensity", "按天求平均"],
            ["派生变量", "sub_metering_remainder", "由总有功功率与分表能耗计算"],
            ["天气变量", "RR, NBJRR1, NBJRR5, NBJRR10, NBJBROU", "按月份合并到每日样本"],
        ],
        widths=[1.25, 3.0, 2.1],
    )

    heading(doc, "2. 模型", level=1)
    p(doc, "设输入窗口长度为 L=90，特征维度为 d，预测长度为 H，其中 H 分别取 90 和 365。所有模型的输入均为 X ∈ R^{L×d}，输出为未来 H 天的 global_active_power 序列。为避免不同预测尺度之间的参数混用，90 天预测与 365 天预测分别训练独立模型。")
    p(doc, "本文采用直接多步预测策略，即模型一次前向传播输出完整的未来序列，而不是逐日递推。该策略可以避免递推预测中误差随步长累积的问题，也使 90 天和 365 天预测具有统一的训练目标。训练阶段最小化预测序列与真实序列之间的均方误差，测试阶段同时报告 MSE 与 MAE；其中 MSE 对较大偏差更敏感，MAE 更直观地反映平均绝对预测误差。")

    heading(doc, "2.1 LSTM 模型", level=2)
    p(doc, "LSTM 通过输入门、遗忘门和输出门控制时间信息的保留与更新，适合处理具有连续依赖关系的时间序列。本文使用两层 LSTM 对 90 天输入窗口进行编码，取最后一层隐藏状态作为窗口表示，再通过 LayerNorm、Dropout 和全连接预测头一次性输出未来 H 天功率曲线。该模型结构相对紧凑，能够作为循环神经网络类方法的基准。", after=1.5)
    add_picture(doc, ASSET_DIR / "model_lstm.png", "图 1  LSTM 预测模型结构示意图", 5.75)
    add_pseudocode(
        doc,
        "算法 1  LSTM 前向预测流程",
        [
            "Input: X in R^{90 x d}, forecast horizon H",
            "Z = LSTM(X)                      # sequence encoding",
            "h = last_hidden_state(Z)         # window representation",
            "y_hat = MLP(LayerNorm(h))        # direct multi-step output",
            "Return y_hat in R^H",
        ],
    )

    heading(doc, "2.2 Transformer 模型", level=2)
    p(doc, "Transformer 使用自注意力机制直接建模输入窗口中任意两个日期之间的依赖关系，避免循环结构逐步传播信息的限制。本文首先将每日多变量特征线性映射到隐空间，并加入位置编码；随后通过多层 Transformer Encoder 提取全局时序依赖，最后使用末时间步表示进行多步预测。与 LSTM 相比，Transformer 更强调全局依赖建模，对输入窗口内不同日期之间的关系具有更直接的表达能力。", after=1.5)
    add_picture(doc, ASSET_DIR / "model_transformer.png", "图 2  Transformer 预测模型结构示意图", 5.75)
    add_pseudocode(
        doc,
        "算法 2  Transformer 前向预测流程",
        [
            "Input: X in R^{90 x d}, forecast horizon H",
            "E = Linear(X) + PositionalEncoding",
            "Z = TransformerEncoder(E)",
            "h = Z[-1]                         # representation of last day",
            "y_hat = MLP(h)",
            "Return y_hat in R^H",
        ],
    )

    heading(doc, "2.3 改进模型：多尺度卷积 Transformer", level=2)
    p(doc, "基础 Transformer 对全局依赖建模能力较强，但局部模式提取能力主要依赖数据学习。家庭用电序列同时包含短期波动、周周期和更长局部趋势，因此本文提出多尺度卷积 Transformer：在 Transformer Encoder 前加入卷积核大小为 3、7、15 的一维卷积分支，分别提取不同时间尺度上的局部特征；再将卷积特征与线性投影特征拼接融合，送入 Transformer Encoder 并通过注意力池化得到窗口表示。该结构旨在兼顾局部周期提取与全局依赖建模，尤其面向 365 天长期预测。", after=1.5)
    p(doc, "多尺度卷积结构的设计动机在于，家庭用电序列中的局部模式并不只对应单一时间尺度。较小卷积核有利于捕获连续几天内的短期波动，较大卷积核则能够覆盖更长的局部变化区间。将不同尺度的局部表示与线性投影分支融合后，再交由自注意力层建模全局关系，可以在模型复杂度可控的前提下增强对局部周期和趋势片段的表达能力。", after=1.5)
    add_picture(doc, ASSET_DIR / "model_msconv_transformer.png", "图 3  多尺度卷积 Transformer 结构示意图", 5.85)
    add_pseudocode(
        doc,
        "算法 3  多尺度卷积 Transformer 前向预测流程",
        [
            "Input: X in R^{90 x d}, forecast horizon H",
            "B0 = Linear(X)",
            "B1, B2, B3 = Conv1D_k=3(X), Conv1D_k=7(X), Conv1D_k=15(X)",
            "E = Fuse(Concat(B0, B1, B2, B3)) + PositionalEncoding",
            "Z = TransformerEncoder(E)",
            "h = AttentionPooling(Z)",
            "y_hat = MLP(h)",
            "Return y_hat in R^H",
        ],
    )

    doc.add_page_break()
    heading(doc, "3. 结果与分析", level=1)
    p(doc, "实验使用过去 90 天作为输入窗口，分别预测未来 90 天和 365 天。所有数值特征仅使用训练集统计量进行标准化，以降低数据泄漏风险。优化器采用 AdamW，学习率为 1e-3，权重衰减为 1e-4；损失函数采用均方误差；训练最大 epoch 为 80，早停 patience 为 12。每个模型与预测长度组合运行 5 个随机种子，并报告 MSE 与 MAE 的均值和标准差。")
    p(doc, "评价时，MSE 反映整体平方误差水平，对峰值预测失败或较大偏差更敏感；MAE 以原始量纲衡量平均绝对误差，更便于比较模型的平均偏离程度。因此，当两个模型在 MSE 与 MAE 上排序不完全一致时，可以理解为二者在峰值刻画和平稳区间拟合之间存在不同权衡。")
    academic_table(
        doc,
        "表 2  实验超参数设置",
        ["项目", "设置"],
        [
            ["输入窗口", "90 天"],
            ["预测长度", "90 天、365 天"],
            ["重复实验", "每个模型与预测长度组合 5 个随机种子"],
            ["评价指标", "MSE、MAE，报告均值与标准差"],
            ["隐藏维度", "128"],
            ["模型层数", "2"],
            ["Dropout", "0.2"],
            ["优化器", "AdamW，学习率 1e-3，权重衰减 1e-4"],
            ["早停策略", "最大 80 epoch，patience = 12"],
        ],
        widths=[1.7, 4.5],
    )
    academic_table(
        doc,
        "表 3  三类模型在不同预测长度上的测试结果",
        ["模型", "预测长度", "MSE 均值 ± 标准差", "MAE 均值 ± 标准差", "验证损失均值", "轮数"],
        summary_rows,
        widths=[1.35, 0.85, 1.45, 1.45, 1.0, 0.55],
    )
    add_picture(doc, ASSET_DIR / "metric_mse.png", "图 4  不同模型在 90 天与 365 天预测任务上的 MSE 对比", 6.15)
    add_picture(doc, ASSET_DIR / "metric_mae.png", "图 5  不同模型在 90 天与 365 天预测任务上的 MAE 对比", 6.15)
    p(doc, "从表 3 和图 4、图 5 可以看出，90 天短期预测中 Transformer 的平均 MSE 最低，LSTM 的平均 MAE 略低，二者整体表现接近；多尺度卷积 Transformer 在短期预测中的误差较高。这说明在相对较短的预测范围内，输入窗口最近阶段的用电水平对目标曲线具有较强解释力，基础 LSTM 与 Transformer 已能学习主要变化模式，而更复杂的多尺度结构可能带来额外参数和优化难度。")
    p(doc, "进一步比较短期预测的两个误差指标可以发现，Transformer 的 MSE 优势说明其在较大偏差样本上的控制更好，而 LSTM 的 MAE 略低则表明其在多数普通日期上的平均偏离较小。这种差异符合两类结构的建模特点：LSTM 的递归状态压缩对平滑序列具有较强归纳偏置，Transformer 则更容易通过注意力权重利用窗口内不同日期之间的非局部关联。")
    p(doc, "在 365 天长期预测中，模型排名发生明显变化。多尺度卷积 Transformer 的 MSE 为 393137.23 ± 44267.92，MAE 为 503.04 ± 35.03，均明显优于 LSTM 和基础 Transformer。该结果说明长期预测不仅需要捕获输入窗口末端的状态，还需要建模更稳定的局部周期和趋势模式。多尺度卷积分支对 3 天、7 天和 15 天局部上下文进行显式抽取，有助于 Transformer 在较长预测范围内获得更稳健的表示。")
    p(doc, "从稳定性看，各模型在 5 个随机种子下的标准差存在差异。短期任务中 LSTM 和 Transformer 的波动相对可控，说明模型在该预测长度下较容易收敛到相近解；长期任务中 Transformer 的 MSE 标准差较大，表明其对初始化和训练过程更敏感。多尺度卷积 Transformer 在长期任务上的均值优势较明显，即使考虑标准差，其整体误差仍低于 LSTM，并与基础 Transformer 拉开差距。")
    doc.add_page_break()
    academic_table(
        doc,
        "表 4  用于绘制代表性预测曲线的实验",
        ["模型", "预测长度", "seed", "MSE", "MAE"],
        best_rows,
        widths=[1.65, 0.9, 0.8, 1.4, 1.2],
    )
    fig_no = 6
    for horizon in [90, 365]:
        for model in MODEL_ORDER:
            seed = int(best[(best["model"] == model) & (best["horizon"] == horizon)]["seed"].iloc[0])
            add_picture(
                doc,
                ASSET_DIR / f"prediction_{model}_{horizon}.png",
                f"图 {fig_no}  {MODEL_LABELS[model]} 在 {horizon} 天预测任务上的预测曲线与真实曲线对比（seed={seed}）",
                6.15,
            )
            fig_no += 1

    heading(doc, "4. 讨论", level=1)
    p(doc, "实验结果表明，不同模型适合的预测尺度并不完全一致。LSTM 通过门控机制对短期序列状态进行压缩，能够在 90 天预测中取得较低 MAE；Transformer 对输入窗口内全局关系建模能力较强，在短期 MSE 指标上表现最好；而多尺度卷积 Transformer 在长期预测上优势明显，说明显式提取局部多尺度模式能够缓解长期预测中的趋势漂移与误差扩散问题。")
    p(doc, "从预测曲线可以看到，三类模型均能学习总体上升或下降趋势，但对尖峰负荷的刻画仍存在不足。该现象一方面与日级聚合后高频信息被压缩有关，另一方面也与直接多步预测方式有关：模型一次性输出未来整段序列，容易倾向于学习平滑趋势，从而低估突发峰值。对于家庭用电场景，峰值往往与设备集中使用、家庭活动和天气突变有关，仅依赖历史用电与月度天气特征仍难以完全解释。")
    p(doc, "本文仍存在一定局限。首先，天气数据为月度粒度，虽然能够反映季节与降水水平，但无法描述日级温度、湿度和突发天气变化；其次，UCI 数据仅来自单户家庭，模型学习到的行为模式可能具有家庭特异性；再次，当前模型采用直接多步预测方式，一次性输出未来所有日期，在极长期预测中仍可能出现过度平滑或峰值刻画不足的问题。")
    p(doc, "后续工作可从三个方向改进：一是加入星期、月份、节假日等日历变量，增强对行为周期的刻画；二是引入更高时间分辨率的天气特征，特别是温度、湿度和体感温度等与空调负荷密切相关的变量；三是尝试分解式预测框架，将序列拆分为趋势项、季节项和残差项后分别建模，以提升长期预测稳定性。")

    centered(doc, "参考文献", 12, True)
    refs = [
        "Dua D, Graff C. UCI Machine Learning Repository[EB/OL]. Irvine, CA: University of California, School of Information and Computer Science, 2019.",
        "Hébrail G, Bérard A. Individual household electric power consumption Data Set[DB/OL]. UCI Machine Learning Repository, 2012.",
        "Météo-France. Données climatologiques de base - mensuelles[DB/OL]. data.gouv.fr.",
        "Hochreiter S, Schmidhuber J. Long short-term memory[J]. Neural Computation, 1997, 9(8): 1735-1780.",
        "Vaswani A, Shazeer N, Parmar N, et al. Attention is all you need[C]//Advances in Neural Information Processing Systems. 2017: 5998-6008.",
        "Kingma D P, Ba J. Adam: A method for stochastic optimization[C]//International Conference on Learning Representations. 2015.",
        "Loshchilov I, Hutter F. Decoupled weight decay regularization[C]//International Conference on Learning Representations. 2019.",
        "Srivastava N, Hinton G, Krizhevsky A, et al. Dropout: A simple way to prevent neural networks from overfitting[J]. Journal of Machine Learning Research, 2014, 15(1): 1929-1958.",
        "Bai S, Kolter J Z, Koltun V. An empirical evaluation of generic convolutional and recurrent networks for sequence modeling[EB/OL]. arXiv:1803.01271, 2018.",
        "Lim B, Zohren S. Time-series forecasting with deep learning: a survey[J]. Philosophical Transactions of the Royal Society A, 2021, 379(2194): 20200209.",
    ]
    for ref in refs:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(ref)
        set_run_font(run, 8.8)

    doc.save(DOCX_PATH)

    md_lines = [
        "# 家庭电力消耗多变量时间序列预测研究",
        "",
        "正文主结构：1.问题介绍、2.模型、3.结果与分析、4.讨论。正式提交版本请使用 PDF。",
        "",
        "## 表 3 三类模型在不同预测长度上的测试结果",
        "",
        markdown_table(["模型", "预测长度", "MSE 均值 ± 标准差", "MAE 均值 ± 标准差", "验证损失均值", "轮数"], summary_rows),
        "",
        "## 表 4 用于绘制代表性预测曲线的实验",
        "",
        markdown_table(["模型", "预测长度", "seed", "MSE", "MAE"], best_rows),
    ]
    MD_PATH.write_text("\n".join(md_lines), encoding="utf-8")
    print(DOCX_PATH)
    print(MD_PATH)


if __name__ == "__main__":
    main()
