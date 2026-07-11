from docx import Document


doc = Document()
doc.add_heading("Render smoke test", level=1)
doc.add_paragraph("This document is used to test LibreOffice headless rendering.")
doc.save("reports/render_smoke.docx")
