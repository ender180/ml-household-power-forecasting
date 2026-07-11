from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


RESULTS_DIR = Path("results_weather")
REPORT_DIR = Path("reports")
ASSET_DIR = REPORT_DIR / "assets"
DOCX_PATH = REPORT_DIR / "机器学习课程考核报告_家庭电力消耗预测.docx"
MD_PATH = REPORT_DIR / "机器学习课程考核报告_家庭电力消耗预测.md"

MODEL_LABELS = {
    "lstm": "LSTM",
    "transformer": "Transformer",
    "conv_transformer": "多尺度卷积 Transformer",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def add_picture_if_exists(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(80, 80, 80)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.22)
    p.add_run(text)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def fmt(mean: float, std: float) -> str:
    return f"{mean:.2f} ± {std:.2f}"


def markdown_table(df: pd.DataFrame) -> str:
    headers = [str(c) for c in df.columns]
    rows = [[str(v) for v in row] for row in df.to_numpy()]
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def build_markdown(summary: pd.DataFrame, best: pd.DataFrame, all_runs: pd.DataFrame) -> str:
    lines = [
        "# 家庭电力消耗多变量时间序列预测实验报告",
        "",
        "作者：请填写姓名、学号、研究方向",
        "",
        "GitHub 链接：请填写代码仓库链接",
        "",
        "## 1. 问题介绍",
        "",
        "本项目面向家庭电力消耗预测任务，目标是利用过去 90 天的多变量用电与天气特征，预测未来 90 天和 365 天的每日总有功功率。该任务可用于家庭用电规划、异常用电检测和智能电网负荷调度。",
        "",
        "原始用电数据来自 UCI Individual household electric power consumption 数据集，天气数据来自 Météo-France 在 data.gouv.fr 发布的月度基础气候数据。本实验将分钟级用电数据按天聚合，并将月度天气特征按月份合并到每日样本。",
        "",
        "## 2. 模型",
        "",
        "本实验比较三类模型：LSTM、Transformer 和多尺度卷积 Transformer。三类模型均分别训练 90 天短期预测模型和 365 天长期预测模型，长期模型参数不复用短期模型参数。",
        "",
        "改进模型在 Transformer Encoder 前加入 3、7、15 三个卷积核大小的一维卷积分支，用于提取短期波动、周尺度模式和更长局部趋势，再通过自注意力建模全局依赖。",
        "",
        "## 3. 结果与分析",
        "",
        markdown_table(summary),
        "",
        "代表性绘图采用每个模型和预测长度中测试 MSE 最低的 seed：",
        "",
        markdown_table(best),
        "",
        "短期 90 天预测中，Transformer 的平均 MSE 最低，LSTM 的 MAE 略优；改进模型由于结构更复杂，在短期任务中没有取得优势。长期 365 天预测中，多尺度卷积 Transformer 的 MSE 和 MAE 均明显优于 LSTM 与基础 Transformer，说明局部多尺度特征提取有助于长期趋势建模。",
        "",
        "训练日志中进度条未达到 80/80 是早停策略导致的正常现象。最大 epoch 为 80，patience 为 12；当验证集损失连续 12 个 epoch 未改善时提前停止，并保存验证集表现最好的模型。",
        "",
        "## 4. 讨论",
        "",
        "实验结果表明，模型结构对不同预测长度的影响不同。短期预测更依赖最近用电状态，基础 Transformer 已能较好捕获时序依赖；长期预测则需要同时处理局部周期、季节性变化和趋势漂移，因此多尺度卷积与 Transformer 的组合更有优势。",
        "",
        "本实验仍有改进空间：天气数据为月度粒度，不能反映日级天气突变；测试集仅来自单个家庭，模型泛化性有限；此外可进一步加入星期、月份、节假日等日历特征，或采用分解式模型分别学习趋势项、周期项和残差项。",
        "",
        "本报告文字整理过程中使用了 ChatGPT/Codex 辅助，但实验代码、数据处理、模型训练和结果分析均基于本项目实际运行结果。",
        "",
        "## 参考文献",
        "",
        "[1] UCI Machine Learning Repository. Individual household electric power consumption dataset.",
        "",
        "[2] Météo-France. Données climatologiques de base - mensuelles. data.gouv.fr.",
        "",
        "[3] Hochreiter, S., and Schmidhuber, J. Long Short-Term Memory. Neural Computation, 1997.",
        "",
        "[4] Vaswani, A. et al. Attention Is All You Need. NeurIPS, 2017.",
    ]
    return "\n".join(lines)


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    summary = pd.read_csv(RESULTS_DIR / "summary.csv")
    all_runs = pd.read_csv(RESULTS_DIR / "all_runs.csv")
    best = pd.read_csv(ASSET_DIR / "best_plot_seeds.csv")

    table_rows = []
    for _, r in summary.sort_values(["horizon", "model"]).iterrows():
        table_rows.append(
            [
                MODEL_LABELS[r["model"]],
                str(int(r["horizon"])),
                fmt(r["mse_mean"], r["mse_std"]),
                fmt(r["mae_mean"], r["mae_std"]),
                str(int(r["runs"])),
            ]
        )

    best_rows = []
    for _, r in best.sort_values(["horizon", "model"]).iterrows():
        best_rows.append(
            [
                MODEL_LABELS[r["model"]],
                str(int(r["horizon"])),
                str(int(r["seed"])),
                f"{r['test_mse']:.2f}",
                f"{r['test_mae']:.2f}",
            ]
        )

    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("家庭电力消耗多变量时间序列预测实验报告")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(8)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("作者：请填写姓名、学号、研究方向    GitHub 链接：请填写代码仓库链接").italic = True

    doc.add_heading("1. 问题介绍", level=1)
    paragraph(
        doc,
        "本项目面向家庭电力消耗预测任务，目标是利用过去 90 天的多变量用电与天气特征，预测未来 90 天和 365 天的每日总有功功率 global_active_power。该任务可用于家庭用电规划、异常用电检测和智能电网负荷调度。",
    )
    paragraph(
        doc,
        "原始用电数据来自 UCI Individual household electric power consumption 数据集，采集自法国一户家庭，时间跨度为 2006 年 12 月至 2010 年 11 月，原始粒度为分钟级。实验按课程要求将分钟级数据聚合为日级数据：功率与分表能耗取日总和，电压和电流取日平均。",
    )
    paragraph(
        doc,
        "外部天气数据来自 Météo-France 在 data.gouv.fr 发布的月度基础气候数据。预处理脚本在巴黎及周边省份气象站中选择覆盖 2006-2010 年且课程指定字段完整的 PARIS-MONTSOURIS 站，并将 RR、NBJRR1、NBJRR5、NBJRR10、NBJBROU 按月份合并到每日样本。",
    )

    doc.add_heading("2. 模型", level=1)
    doc.add_heading("2.1 LSTM", level=2)
    paragraph(
        doc,
        "LSTM 使用过去 90 天的多变量序列作为输入，通过循环门控结构建模时间依赖关系。模型取最后一层隐藏状态，经 LayerNorm、Dropout 和全连接预测头一次性输出未来 H 天曲线，其中 H 分别为 90 和 365。",
    )
    doc.add_heading("2.2 Transformer", level=2)
    paragraph(
        doc,
        "Transformer 将每日特征投影到隐空间，加入正弦位置编码后送入多层 Transformer Encoder。自注意力机制用于捕获不同日期之间的全局依赖，最后一个时间步的表示用于预测未来曲线。",
    )
    doc.add_heading("2.3 改进模型：多尺度卷积 Transformer", level=2)
    paragraph(
        doc,
        "改进模型在 Transformer Encoder 前加入多尺度一维卷积分支，卷积核大小分别为 3、7、15，用于提取短期波动、周尺度模式和更长跨度的局部趋势。卷积特征与线性投影特征拼接后送入 Transformer Encoder，再通过注意力池化得到序列表示并输出预测曲线。该结构的动机是先增强局部模式抽取能力，再利用自注意力建模长距离依赖。",
    )

    doc.add_heading("3. 实验设置", level=1)
    bullet(doc, "输入窗口长度为 90 天，预测长度分别为 90 天和 365 天。")
    bullet(doc, "三类模型在两种预测长度上分别训练，长期预测模型参数不复用短期预测模型参数。")
    bullet(doc, "每个模型和预测长度组合使用 5 个随机种子，报告 MSE 和 MAE 的均值与标准差。")
    bullet(doc, "训练最大 epoch 为 80，采用验证集早停策略，patience 为 12，并保存验证集表现最好的模型。")

    doc.add_heading("4. 结果与分析", level=1)
    add_table(doc, ["模型", "预测长度", "MSE 均值 ± std", "MAE 均值 ± std", "轮数"], table_rows)
    add_picture_if_exists(doc, ASSET_DIR / "mse_comparison.png", "图 1  三类模型在 90 天与 365 天预测任务上的 MSE 对比")
    add_picture_if_exists(doc, ASSET_DIR / "mae_comparison.png", "图 2  三类模型在 90 天与 365 天预测任务上的 MAE 对比")
    paragraph(
        doc,
        "从表格和对比图可以看出，90 天短期预测中 Transformer 的平均 MSE 最低，LSTM 的平均 MAE 略低；改进模型在短期预测中没有取得优势，说明短期任务更依赖最近用电状态，基础时序模型已经能较好捕捉主要变化。",
    )
    paragraph(
        doc,
        "365 天长期预测中，多尺度卷积 Transformer 的 MSE 和 MAE 均明显低于 LSTM 与基础 Transformer，说明局部多尺度模式提取对长期趋势和周期性变化建模有帮助。这也符合改进模型的设计动机：卷积分支负责抽取不同时间尺度的局部特征，Transformer 负责整合全局依赖。",
    )

    doc.add_heading("4.1 代表性预测曲线", level=2)
    add_table(doc, ["模型", "预测长度", "绘图 seed", "MSE", "MAE"], best_rows)
    for horizon in [90, 365]:
        for model in ["lstm", "transformer", "conv_transformer"]:
            seed = int(best[(best["model"] == model) & (best["horizon"] == horizon)]["seed"].iloc[0])
            img = RESULTS_DIR / model / f"horizon_{horizon}" / f"seed_{seed}" / "prediction_vs_ground_truth.png"
            add_picture_if_exists(
                doc,
                img,
                f"图  {MODEL_LABELS[model]} 在 {horizon} 天预测任务上的预测曲线与真实曲线对比（seed={seed}）",
                width=6.1,
            )

    doc.add_heading("5. 讨论", level=1)
    paragraph(
        doc,
        "训练日志中进度条没有达到 80/80 是早停机制导致的正常现象。代码设置最大 epoch 为 80、patience 为 12，当验证集损失连续 12 个 epoch 没有改善时提前停止，并保存验证集表现最好的参数。所有 30 次实验均已产生 metrics.json、history.csv 和 prediction_vs_ground_truth.png，说明训练过程完整完成。",
    )
    paragraph(
        doc,
        "本实验的主要限制在于天气数据为月度粒度，不能反映日级天气突变；数据来自单个家庭，模型泛化性有限；此外家庭用电受到成员行为、节假日和设备使用习惯影响，仅依赖用电变量和月度天气变量仍难完全解释长期波动。",
    )
    paragraph(
        doc,
        "后续可进一步加入星期、月份、节假日等日历特征，或采用分解式模型分别学习趋势项、周期项和残差项。也可以尝试 ProbSparse Transformer、Informer、N-BEATS 等结构，提升长期预测稳定性。",
    )

    doc.add_heading("6. 作者贡献与工具说明", level=1)
    paragraph(doc, "作者贡献：请根据实际组队情况填写各位作者姓名、所属研究方向和具体贡献。")
    paragraph(
        doc,
        "本报告文字整理过程中使用了 ChatGPT/Codex 辅助，但实验代码、数据处理、模型训练和结果分析均基于本项目实际运行结果。提交前请将作者信息和 GitHub 链接替换为真实内容。",
    )

    doc.add_heading("参考文献", level=1)
    refs = [
        "UCI Machine Learning Repository. Individual household electric power consumption dataset.",
        "Météo-France. Données climatologiques de base - mensuelles. data.gouv.fr.",
        "Hochreiter, S., and Schmidhuber, J. Long Short-Term Memory. Neural Computation, 1997.",
        "Vaswani, A. et al. Attention Is All You Need. NeurIPS, 2017.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    doc.save(DOCX_PATH)
    md_summary = summary.copy()
    md_summary["model"] = md_summary["model"].map(MODEL_LABELS)
    md_summary["MSE 均值 ± std"] = md_summary.apply(lambda r: fmt(r["mse_mean"], r["mse_std"]), axis=1)
    md_summary["MAE 均值 ± std"] = md_summary.apply(lambda r: fmt(r["mae_mean"], r["mae_std"]), axis=1)
    md_summary = md_summary[["model", "horizon", "MSE 均值 ± std", "MAE 均值 ± std", "runs"]]
    best_md = best.copy()
    best_md["model"] = best_md["model"].map(MODEL_LABELS)
    MD_PATH.write_text(build_markdown(md_summary, best_md, all_runs), encoding="utf-8")
    print(DOCX_PATH)
    print(MD_PATH)


if __name__ == "__main__":
    main()
